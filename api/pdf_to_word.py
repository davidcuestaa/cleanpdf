"""
API endpoint: POST /api/pdf_to_word
Convierte PDF a DOCX.

Entorno: Vercel serverless.
Librería principal: pdf2docx  (puro Python, funciona en Vercel)
Fallback: pdfminer.six + python-docx (solo texto estructurado)

pdf2docx reconstruye layout, tablas e imágenes del PDF.
Instalar: pip install pdf2docx
"""
import os
import io
import cgi
import json
import tempfile
import traceback
from http.server import BaseHTTPRequestHandler

try:
    from pdf2docx import Converter
    PDF2DOCX_OK = True
except ImportError:
    PDF2DOCX_OK = False

try:
    from pdfminer.high_level import extract_pages
    from pdfminer.layout import LTTextContainer, LTChar
    PDFMINER_OK = True
except ImportError:
    PDFMINER_OK = False

try:
    from docx import Document as DocxDoc
    from docx.shared import Pt, Inches
    DOCX_OK = True
except ImportError:
    DOCX_OK = False


def _parse_multipart(handler):
    ct = handler.headers.get("Content-Type", "")
    if "multipart/form-data" not in ct:
        return None, None, "Content-Type must be multipart/form-data"
    env = {"REQUEST_METHOD": "POST", "CONTENT_TYPE": ct,
           "CONTENT_LENGTH": handler.headers.get("Content-Length", "0")}
    body = handler.rfile.read(int(handler.headers.get("Content-Length", 0)))
    fs = cgi.FieldStorage(fp=io.BytesIO(body), environ=env, keep_blank_values=True)
    if "file" not in fs:
        return None, None, "No 'file' field in request"
    f = fs["file"]
    return f.file.read(), f.filename, None


def convert_pdf2docx(file_bytes):
    """Conversión de alta calidad con pdf2docx."""
    with tempfile.TemporaryDirectory() as tmp:
        pdf_path  = os.path.join(tmp, "input.pdf")
        docx_path = os.path.join(tmp, "output.docx")

        with open(pdf_path, "wb") as f:
            f.write(file_bytes)

        cv = Converter(pdf_path)
        cv.convert(
            docx_path,
            start=0,
            end=None,
            # Tolerancias ajustadas para mejor reconstrucción de layout
            connected_border_tolerance=0.5,
            line_overlap_threshold=0.9,
            line_break_free_space_ratio=0.1,
            lines_left_aligned_threshold=0.1,
            lines_right_aligned_threshold=0.1,
            debug=False,
        )
        cv.close()

        with open(docx_path, "rb") as f:
            return f.read()


def convert_fallback(file_bytes):
    """Fallback texto-estructurado con pdfminer + python-docx."""
    if not (PDFMINER_OK and DOCX_OK):
        raise RuntimeError(
            "Instala pdf2docx: pip install pdf2docx\n"
            "O alternativamente: pip install pdfminer.six python-docx"
        )

    doc = DocxDoc()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    for page_layout in extract_pages(io.BytesIO(file_bytes)):
        elements = sorted(
            [e for e in page_layout if isinstance(e, LTTextContainer)],
            key=lambda e: -e.y1,
        )
        for element in elements:
            text = element.get_text().strip()
            if not text:
                continue

            # Detectar tamaño y negrita del primer carácter
            font_size = 11.0
            is_bold   = False
            for line in element:
                for char in line:
                    if isinstance(char, LTChar):
                        font_size = round(char.size, 1)
                        is_bold   = "Bold" in (char.fontname or "")
                        break
                break

            para = doc.add_paragraph()
            run  = para.add_run(text)
            run.font.size = Pt(font_size)
            run.font.bold = is_bold or font_size >= 14

        doc.add_page_break()

    # Quitar último salto de página
    try:
        last = doc.paragraphs[-1]
        p = last._element
        p.getparent().remove(p)
    except Exception:
        pass

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class handler(BaseHTTPRequestHandler):

    def do_OPTIONS(self):
        self._cors()
        self.end_headers()

    def do_POST(self):
        try:
            file_bytes, filename, err = _parse_multipart(self)
            if err:
                self._error(400, err)
                return
            if not filename or not filename.lower().endswith(".pdf"):
                self._error(400, "Por favor sube un archivo PDF (.pdf)")
                return
            if not file_bytes:
                self._error(400, "El archivo está vacío")
                return

            if PDF2DOCX_OK:
                try:
                    docx_bytes = convert_pdf2docx(file_bytes)
                except Exception:
                    docx_bytes = convert_fallback(file_bytes)
            else:
                docx_bytes = convert_fallback(file_bytes)

            out_name = filename.rsplit(".", 1)[0] + ".docx"
            self.send_response(200)
            self.send_header(
                "Content-Type",
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document",
            )
            self.send_header("Content-Disposition",
                             f'attachment; filename="{out_name}"')
            self.send_header("Content-Length", str(len(docx_bytes)))
            self._cors()
            self.end_headers()
            self.wfile.write(docx_bytes)

        except Exception as e:
            self._error(500, f"Error al convertir: {str(e).splitlines()[0]}")

    def _cors(self):
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")

    def _error(self, code, msg):
        body = json.dumps({"error": msg}).encode()
        self.send_response(code)
        self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", str(len(body)))
        self._cors()
        self.end_headers()
        self.wfile.write(body)

    def log_message(self, fmt, *args):
        pass
